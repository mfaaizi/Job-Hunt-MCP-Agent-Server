"""
Raw text extraction from resume files (PDF / DOCX).
This is the only module that touches the original file — everything downstream
works off the extracted text or the normalized JSON.
"""
import re
from pathlib import Path

import pdfplumber
from docx import Document

# Icon fonts (e.g. FontAwesome glyphs used for contact-info icons like  /  / )
# often extract as unmapped character codes like "(cid:239)". These carry no
# semantic meaning, so strip them rather than passing noise to the LLM.
_CID_ARTIFACT_RE = re.compile(r"\(cid:\d+\)\s*")

# Fallback patterns, used only if no hyperlink annotations are found at all
# (e.g. a resume where contact info is plain text, not clickable links).
_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"\+?\d[\d\s().-]{7,}\d")


def extract_contact_links_from_pdf(file_path: Path) -> dict:
    """
    Pulls contact info directly from PDF hyperlink annotations (mailto:,
    tel:, github.com/..., linkedin.com/...). This is far more reliable than
    parsing visible text: resumes often render contact info as an icon glyph
    next to a label, and the icon extracts as garbage — but the underlying
    link target (the actual URI behind the clickable text) is unaffected by
    that, since it's stored as PDF metadata, not rendered glyphs.

    Returns a dict with whichever of email/phone/github/linkedin/portfolio
    it found. Empty dict if the PDF has no hyperlink annotations at all.
    """
    result: dict = {}

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            for hyperlink in getattr(page, "hyperlinks", []):
                uri = hyperlink.get("uri", "")
                if not uri:
                    continue

                if uri.startswith("mailto:"):
                    result.setdefault("email", uri[len("mailto:"):])
                elif uri.startswith("tel:"):
                    result.setdefault("phone", uri[len("tel:"):])
                elif "github.com" in uri.lower():
                    result.setdefault("github", uri)
                elif "linkedin.com" in uri.lower():
                    result.setdefault("linkedin", uri)
                else:
                    # Any other external link near the header is most likely
                    # a personal site/portfolio link.
                    result.setdefault("portfolio", uri)

    return result


def extract_name_fallback(raw_text: str) -> str | None:
    """
    Resumes near-universally put the person's name as the very first
    non-empty line. Used only when the LLM returns null for contact.name —
    a simple positional heuristic beats asking a small model to identify
    "which line is the name" from noisy header text.
    """
    for line in raw_text.splitlines():
        line = line.strip()
        if not line:
            continue
        # Skip lines that are clearly not a name (too long, looks like a
        # sentence/section header, contains contact-info markers).
        if len(line) > 60 or "@" in line or any(ch.isdigit() for ch in line):
            return None
        return line
    return None


def extract_contact_fields_fallback(raw_text: str) -> dict:
    """
    Regex-based fallback for email/phone when there are no PDF hyperlinks to
    pull from (e.g. DOCX files, or a PDF where contact info isn't a clickable
    link). Less reliable than extract_contact_links_from_pdf, so callers
    should prefer that first and only use this as a backstop.
    """
    result: dict = {}

    email_match = _EMAIL_RE.search(raw_text)
    if email_match:
        result["email"] = email_match.group(0)

    phone_match = _PHONE_RE.search(raw_text)
    if phone_match:
        result["phone"] = re.sub(r"\s+", " ", phone_match.group(0)).strip()

    return result


def _clean_extracted_text(text: str) -> str:
    return _CID_ARTIFACT_RE.sub("", text)


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract plain text from a PDF, page by page, preserving line breaks."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract plain text from a DOCX, paragraph by paragraph (including tables)."""
    doc = Document(file_path)
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # DOCX resumes sometimes use tables for layout (e.g. skills grids) — capture those too
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def extract_resume_text(file_path: str | Path) -> str:
    """
    Dispatch based on file extension. Raises ValueError for unsupported types.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if not file_path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    if suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported resume file type: {suffix} (expected .pdf or .docx)")

    if not text.strip():
        raise ValueError(f"No extractable text found in {file_path}. It may be a scanned/image-based document.")

    return _clean_extracted_text(text)
